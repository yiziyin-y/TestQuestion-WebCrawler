# -*- coding: utf-8 -*- 

import sys
import win32com
import MySQLdb
import win32com.client
from docx import Document
from cStringIO import StringIO
import win32clipboard
from PIL import Image
import ast
from BS import glo

reload(sys)
sys.setdefaultencoding('utf-8')


def getDbConn():
    conn = MySQLdb.connect(
        host='127.0.0.1',
        port=3306,
        user='root',
        passwd='a123',
        db='bs',
        charset='utf8'
    )
    return conn


def export():
    conn = getDbConn()
    cursor = conn.cursor()
    inputType = glo.inputType
    subject = inputType
    print(subject)

    sqlE = "SELECT `ques`,`answer`,`analyze`,`awImage` FROM" + " " + subject
    print("sqlE:" + sqlE)
    cursor.execute(sqlE)
    rowss = cursor.fetchall()

    document = Document()
    i = 1
    for rows in rowss:

        title = "第" + str(i) + "题:"
        print(unicode(title))
        i = i + 1
        document.add_paragraph(unicode(title), style='Normal')
        # 每一题的输出
        for row in rows:
            row1 = row.strip()
            document.add_paragraph(unicode(row1), style='Normal')
            print("ROW:" + row1)

    document.add_page_break()
    # 按科目名字导出DocX
    document.save(subject + '.docx')
    cursor.close()
    conn.close()


def setImageToClipboard(clip_type, data):
    win32clipboard.OpenClipboard()
    win32clipboard.EmptyClipboard()
    win32clipboard.SetClipboardData(clip_type, data)
    win32clipboard.CloseClipboard()


def getImageFromClipboard():
    win32clipboard.OpenClipboard()
    d = win32clipboard.GetClipboardData(win32clipboard.CF_DIB)
    win32clipboard.CloseClipboard()
    return d


def replaceQ():
    conn = getDbConn()
    cursor = conn.cursor()
    inputType = glo.inputType
    subject = inputType
    print(subject)

    sqlImageQ = "SELECT `quesImage` FROM" + " " + subject
    print("sqlImageQ:" + sqlImageQ)
    cursor.execute(sqlImageQ)
    ImagesQ = cursor.fetchall()

    #  循环取出链接元组
    for Image1Q in ImagesQ:

        # 取出每一题的所有图片
        for ImageQ in Image1Q:
            # 判断是否不含图片
            if ImageQ.strip() != "":
                #  将数据库取出的str识别为列表
                print("imageq" + ImageQ)
                listQs = ast.literal_eval(ImageQ)
                # 循环取出列表里的多个图片
                for ImageQS in listQs:

                    print("imageQS:" + ImageQS)
                    ImageQFileUrl = ImageQS.strip()
                    ImageQFileName = ImageQFileUrl[38:]
                    filepath = r'D:\IntelliJ_IDEA2018.1.6\workspace\BS\queImages'

                    if ImageQ.strip() != '':
                        tail = ImageQFileUrl[31:]

                        print("ImageQfilename:" + ImageQFileName)
                        image = Image.open(filepath + '\\' + ImageQFileName)
                        output = StringIO()
                        image.convert("RGB").save(output, "BMP")
                        data = output.getvalue()[14:]
                        output.close()
                        setImageToClipboard(win32clipboard.CF_DIB, data)
                        content = getImageFromClipboard()

                        w = win32com.client.Dispatch('Word.Application')
                        # 后台运行，不显示，不警告
                        # w.Visible = 0
                        # w.DisplayAlerts = 0
                        doc = w.Documents.Open(
                            "D:/IntelliJ_IDEA2018.1.6/workspace/BS/BS/" + subject + ".docx")

                        search = u"https://tiku.21cnjy.com/tikupic" + tail

                        w.Selection.Find.Replacement.ClearFormatting()
                        w.Selection.Find.Execute(
                            search,
                            False,
                            True,
                            False,
                            False,
                            False,
                            True,
                            1,
                            True,
                            "^c",
                            2)

                        doc.SaveAs(
                            "D:/IntelliJ_IDEA2018.1.6/workspace/BS/BS/" +
                            subject +
                            ".docx")
                        doc.Close()

                        w.Quit()


def replaceA():
    conn = getDbConn()
    cursor = conn.cursor()
    inputType = glo.inputType
    subject = inputType
    print(subject)

    sqlImageA = "SELECT `awImage` FROM" + " " + subject
    print("sqlImageA:" + sqlImageA)
    cursor.execute(sqlImageA)
    ImagesA = cursor.fetchall()

    for Image1A in ImagesA:

        for ImageA in Image1A:
            ImageAFileUrl = ImageA.strip()
            ImageAFileName = ImageAFileUrl[43:]
            filepathA = r'D:\IntelliJ_IDEA2018.1.6\workspace\BS\awImages'

            if ImageA.strip() != '':
                tailA = ImageAFileUrl[36:]
                print("tail:" + tailA)
                print("ImageAfilename:" + ImageAFileName)
                image = Image.open(filepathA + '\\' + ImageAFileName)
                output = StringIO()
                image.convert("RGB").save(output, "BMP")
                data = output.getvalue()[14:]
                output.close()
                setImageToClipboard(win32clipboard.CF_DIB, data)
                content = getImageFromClipboard()

                w = win32com.client.Dispatch('Word.Application')
                # 后台运行，不显示，不警告
                # w.Visible = 0
                # w.DisplayAlerts = 0
                doc = w.Documents.Open(
                    "D:/IntelliJ_IDEA2018.1.6/workspace/BS/BS/" + subject + ".docx")

                search = u"https://tiku.21cnjy.com/tikupic" + tailA

                w.Selection.Find.Replacement.ClearFormatting()
                w.Selection.Find.Execute(
                    search,
                    False,
                    True,
                    False,
                    False,
                    False,
                    True,
                    1,
                    True,
                    "^c",
                    2)

                doc.SaveAs(
                    "D:/IntelliJ_IDEA2018.1.6/workspace/BS/BS/" +
                    subject +
                    ".docx")
                doc.Close()

                w.Quit()

#
# if __name__ == "__main__":
#     # export()
#     # replaceQ()
#     replaceA()
